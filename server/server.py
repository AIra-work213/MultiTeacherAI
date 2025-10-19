from fastapi import FastAPI, UploadFile, File, Form, Request
from pydantic import BaseModel
from typing import Optional
import uvicorn
from database.chroma_db import ChromaDB
from database.postgres import init_db, insert_metadata, find_metadata_by_id, find_max_topic_id, find_topic_id
import asyncio
import json
from PyPDF2 import PdfReader
import docx
from models.gigachatModel import GigaModel
from langchain_text_splitters import RecursiveCharacterTextSplitter
from contextlib import asynccontextmanager

model = GigaModel()
vectorstore = ChromaDB()
@asynccontextmanager
async def lifespan(app: FastAPI):
    await init_db()
    app.state.topic_id = (await find_max_topic_id()) or 0
    yield

app = FastAPI(lifespan=lifespan)


class QuestRequest(BaseModel):
    quest: str
    user_id: int

class TestData(BaseModel):
    is_success: bool
    questions: list[str]
    answers: list[str]
class GetTopics(BaseModel):
    user_id: int

class ForTest(BaseModel):
    user_id: int
    topic: str
    topic_ID: Optional[int] = None


@app.post("/upload")
async def upload_file(request: Request, metadata: str = Form(...), file: UploadFile = File(...)):

    metadata_dict = json.loads(metadata)
    myFile = ""

    is_pdf = (file.content_type == 'application/pdf') or (file.filename.endswith('.pdf'))
    is_txt = (file.content_type == 'text/plain') or (file.filename.endswith('.txt'))
    is_docx = (file.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document') or (file.filename.endswith('.docx'))

    if is_pdf:
        metadata_dict["file_type"] = "pdf"
        pdf_read = PdfReader(file.file)
        pdf_text = ""
        try:
            for page in pdf_read.pages:
                pdf_text += page.extract_text()
        except:
            pass
        myFile = pdf_text
    elif is_txt:
        myFile = (await file.read()).decode('utf-8')
    elif is_docx:
        myFile = docx.Document(file.file)
        myFile = docx.Document(file.file)
        docx_text = [p.text for p in myFile.paragraphs]
        myFile = '\n'.join(docx_text)
    else:
        return {"status": "Неподдерживаемый формат файла"}
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=2048,
        chunk_overlap=200,
        length_function=len
    )
    

    metadata_dict["topic"] = model.get_topic(myFile)
    request.app.state.topic_id = getattr(request.app.state, "topic_id", 0) + 1
    metadata_dict["topic_id"] = request.app.state.topic_id

    text_chunks = text_splitter.split_text(myFile)
    metadatas_for_chunks = [
        {'uploader_id': str(metadata_dict["user_id"]), 'file_name': metadata_dict["file_name"], 'topic_id': metadata_dict["topic_id"]}
    ] * len(text_chunks)

    vectorstore.add_texts(texts=text_chunks, metadatas=metadatas_for_chunks)

    await insert_metadata(metadata_dict)
    return {"status": True, "file_name": metadata_dict["file_name"], "topic": metadata_dict["topic"], "topic_id": metadata_dict["topic_id"]}


@app.post("/quest") 
async def quest(request: QuestRequest):
    '''Обрабатывает текстовый запрос пользователя и возвращает ответ'''
    data = vectorstore.query(user_id=request.user_id, quest=request.quest) 
    ans = model.question(request.quest, data)
    return {"answer": ans}


@app.post("/test") # не работает
async def test(data: ForTest):
    '''Проводит тестирование для пользователя по теме'''
    if data.topic_ID is not None:
        id = data.topic_ID
    else:
        res = await find_topic_id(data.topic, data.user_id)
        id = res['topic_id']
    
    all_text = vectorstore.get_texts_by_topic_id(data.user_id, id)
    ans = model.get_questions(all_text)
    questions = ans['questions']
    answers = ans['answers']
    return {"is_success": True, "questions": questions, "answers": answers}


@app.post("/topics") 
async def topics(data: GetTopics):
    topics = await find_metadata_by_id(data.user_id)
    '''Выполняет поиск в SQL тем файлов, который были загружены этим айди (не более 10)'''
    ans = ''
    for t in topics:
        ans += f"Тема: {t['topic']}. ID темы: {t['topic_id']}.\n\n"
    return {"response": ans}

@app.get('/connect')
async def chech_connection():
    return {"status": "Server is running"}


if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)